#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generates DATA_FLOW_AND_ONBOARDING.docx from
DATA_FLOW_AND_ONBOARDING.md (single source of truth: the md).
Deterministic layout; re-run after editing the md. Requires
python-docx. Reference documentation only — no program authority."""
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "DATA_FLOW_AND_ONBOARDING.md")
OUT = os.path.join(HERE, "DATA_FLOW_AND_ONBOARDING.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def add_md_runs(par, text):
    """Bold (**x**) and code (`x`) inline runs."""
    for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            par.add_run(tok)


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j].paragraphs[0]
        add_md_runs(c, h)
        for r in c.runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            if j < len(t.rows[i].cells):
                add_md_runs(t.rows[i].cells[j].paragraphs[0], val)
    return t


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.8)

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            rows = [[c.strip() for c in r.strip().strip("|").split("|")]
                    for r in table_buf]
            body = [r for r in rows[1:]
                    if not all(set(c) <= {"-", ":", " "} for c in r)]
            add_table(doc, rows[0], body)
            doc.add_paragraph()
            table_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_buf))
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Inches(0.3)
            code_buf = []

    for ln in lines:
        if ln.strip().startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(ln)
            continue
        if ln.strip().startswith("|"):
            table_buf.append(ln)
            continue
        flush_table()
        s = ln.strip()
        if not s or s == "---":
            continue
        if s.startswith("# "):
            h = doc.add_heading("", level=0)
            add_md_runs(h, s[2:])
            for r in h.runs:
                r.font.color.rgb = ACCENT
        elif s.startswith("## "):
            h = doc.add_heading("", level=1)
            add_md_runs(h, s[3:])
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            add_md_runs(p, re.sub(r"^\d+\.\s", "", s))
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_md_runs(p, s[2:])
        else:
            p = doc.add_paragraph()
            add_md_runs(p, s)
        i += 1
    flush_table()
    flush_code()

    doc.save(OUT)
    print(f"wrote {os.path.relpath(OUT, HERE)}")


if __name__ == "__main__":
    main()
